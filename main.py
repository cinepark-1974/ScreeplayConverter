rom __future__ import annotations

import os
from datetime import datetime
from io import BytesIO
from typing import Dict

import streamlit as st
from anthropic import Anthropic

from prompt import (
    SYSTEM_PROMPT,
    build_gap_diagnosis_prompt,
    build_intake_merge_prompt,
    build_rewrite_prompt,
    build_story_reinforcement_prompt,
    build_unit_draft_prompt,
    build_unit_plan_prompt,
)

APP_TITLE = "BLUE JEANS NOVEL ENGINE"
ANTHROPIC_MODEL = "claude-sonnet-4-6"
DEFAULT_STYLE = "시드니 셀던 스타일의 대중 장편소설 감각. 빠르게 읽히고, 장면이 선명하며, 챕터 말미 후킹이 강한 문체."

st.set_page_config(
    page_title=APP_TITLE,
    page_icon="👖",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ─────────────────────────────────────
# State
# ─────────────────────────────────────
DEFAULT_STATE = {
    "title": "",
    "genre": "스릴러",
    "style_note": DEFAULT_STYLE,
    "overview": "",
    "characters": "",
    "synopsis": "",
    "extra_notes": "",
    "merged_summary": "",
    "gap_report": "",
    "reinforced_story": "",
    "unit_plan": "",
    "unit_drafts": {},
    "selected_unit": 1,
    "rewrite_mode": "더 상업적으로",
}

for key, value in DEFAULT_STATE.items():
    if key not in st.session_state:
        st.session_state[key] = value


# ─────────────────────────────────────
# CSS
# ─────────────────────────────────────
st.markdown(
    """
<style>
@import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
@import url('https://cdn.jsdelivr.net/gh/projectnoonnu/2408-3@latest/Paperlogy.css');
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700;900&display=swap');

:root {
    --navy: #202A78;
    --y: #FFCB05;
    --bg: #F7F7F5;
    --card: #FFFFFF;
    --card-border: #DDDDE6;
    --t: #2A2A3A;
    --g: #2EC484;
    --dim: #8A8FA3;
    --light-bg: #EEEEF6;
    --display: 'Playfair Display', 'Paperlogy', 'Georgia', serif;
    --body: 'Pretendard', -apple-system, BlinkMacSystemFont, sans-serif;
    --heading: 'Paperlogy', 'Pretendard', sans-serif;
}

html, body, [class*="css"] {
    font-family: var(--body);
    color: var(--t);
    -webkit-font-smoothing: antialiased;
}

.stApp, [data-testid="stAppViewContainer"], [data-testid="stMain"],
[data-testid="stMainBlockContainer"], [data-testid="stHeader"],
[data-testid="stBottom"] {
    background-color: var(--bg) !important;
    color: var(--t) !important;
}

.stMarkdown, .stText, .stCode { color: var(--t) !important; }
h1, h2, h3, h4, h5, h6 {
    color: var(--navy) !important;
    font-family: var(--heading) !important;
}
p, span, label, div, li { color: var(--t); }
section[data-testid="stSidebar"] { display: none; }

.stTextInput input, .stTextArea textarea,
[data-testid="stTextInput"] input, [data-testid="stTextArea"] textarea {
    background-color: var(--card) !important;
    color: var(--t) !important;
    border: 1.5px solid var(--card-border) !important;
    border-radius: 8px !important;
    font-family: var(--body) !important;
    font-size: 0.93rem !important;
    padding: 0.7rem 0.9rem !important;
}

.stTextInput input:focus, .stTextArea textarea:focus,
[data-testid="stTextInput"] input:focus, [data-testid="stTextArea"] textarea:focus {
    border-color: var(--navy) !important;
    box-shadow: 0 0 0 2px rgba(32,42,120,0.08) !important;
}

.stTextInput input::placeholder, .stTextArea textarea::placeholder,
[data-testid="stTextInput"] input::placeholder, [data-testid="stTextArea"] textarea::placeholder {
    color: var(--dim) !important;
    font-size: 0.86rem !important;
}

.stSelectbox > div > div, [data-baseweb="select"] > div, [data-baseweb="select"] input {
    background-color: var(--card) !important;
    color: var(--t) !important;
    border-color: var(--card-border) !important;
    border-radius: 8px !important;
}

[data-baseweb="popover"], [data-baseweb="menu"], [role="listbox"], [role="option"] {
    background-color: var(--card) !important;
    color: var(--t) !important;
}
[role="option"]:hover { background-color: var(--light-bg) !important; }

.stTextInput label, .stTextArea label, .stSelectbox label {
    color: var(--t) !important;
    font-weight: 700 !important;
    font-size: 0.88rem !important;
    margin-bottom: 0.35rem !important;
}

.stButton > button {
    color: var(--t) !important;
    border: 1.5px solid var(--card-border) !important;
    background-color: var(--card) !important;
    border-radius: 8px !important;
    font-family: var(--body) !important;
    font-weight: 800 !important;
    font-size: 0.92rem !important;
    padding: 0.64rem 1.2rem !important;
    transition: all 0.2s;
}

.stButton > button:hover {
    border-color: var(--navy) !important;
    box-shadow: 0 2px 8px rgba(32,42,120,0.08) !important;
}

.stButton > button[kind="primary"],
.stButton > button[data-testid="stBaseButton-primary"] {
    background-color: var(--y) !important;
    color: var(--navy) !important;
    border-color: var(--y) !important;
    font-weight: 900 !important;
}

.stButton > button[kind="primary"]:hover,
.stButton > button[data-testid="stBaseButton-primary"]:hover {
    background-color: #E8B800 !important;
    box-shadow: 0 2px 12px rgba(255,203,5,0.30) !important;
}

.stDownloadButton > button {
    color: var(--navy) !important;
    border: 1.5px solid var(--y) !important;
    background-color: var(--y) !important;
    border-radius: 8px !important;
    font-family: var(--body) !important;
    font-weight: 900 !important;
    font-size: 0.90rem !important;
    padding: 0.64rem 1.2rem !important;
}

.stExpander, details, details summary {
    background-color: var(--card) !important;
    color: var(--t) !important;
    border: 1px solid var(--card-border) !important;
    border-radius: 8px !important;
}

details[open] > div { background-color: var(--card) !important; }
.stExpander summary, .stExpander summary span { color: var(--t) !important; }
.stAlert { color: var(--t) !important; border-radius: 8px !important; }
[data-testid="stVerticalBlock"], [data-testid="stHorizontalBlock"], [data-testid="stColumn"] {
    background-color: transparent !important;
}

.header {
    font-size: 0.85rem;
    font-weight: 700;
    color: var(--navy);
    letter-spacing: 0.15em;
    font-family: var(--heading);
}

.brand-title {
    font-size: 2.6rem;
    font-weight: 900;
    color: var(--navy);
    font-family: var(--display);
    letter-spacing: -0.02em;
    position: relative;
    display: inline-block;
}

.brand-title::after {
    content: '';
    position: absolute;
    bottom: 2px;
    left: 0;
    width: 100%;
    height: 4px;
    background: var(--y);
    border-radius: 2px;
}

.sub {
    font-size: 0.72rem;
    color: var(--dim);
    letter-spacing: 0.15em;
    margin-top: 0.5rem;
    margin-bottom: 1.5rem;
}

.callout {
    background: var(--light-bg);
    border-left: 4px solid var(--navy);
    padding: 0.95rem 1.1rem;
    margin: 0.5rem 0 1rem 0;
    border-radius: 0 8px 8px 0;
    font-size: 0.90rem;
    color: var(--t);
}

.cl {
    color: var(--navy);
    font-weight: 800;
    font-size: 0.74rem;
    letter-spacing: 0.03em;
    margin-bottom: 0.3rem;
    text-transform: uppercase;
}

.section-header {
    background: var(--y);
    color: var(--navy);
    padding: 0.75rem 1rem;
    border-radius: 6px;
    font-weight: 900;
    font-size: 1.04rem;
    font-family: var(--heading);
    margin: 1.6rem 0 0.9rem 0;
    display: flex;
    justify-content: space-between;
    align-items: center;
}

.section-header .en {
    font-family: var(--display);
    font-size: 0.78rem;
    font-weight: 700;
    letter-spacing: 0.05em;
    opacity: 0.72;
}

.small-meta {
    font-size: 0.80rem;
    color: var(--dim);
    margin-top: -0.1rem;
    margin-bottom: 0.55rem;
}

.output-card {
    background: var(--card);
    border: 1px solid var(--card-border);
    border-radius: 10px;
    padding: 1rem 1.1rem;
    margin: 0.55rem 0 1rem 0;
}

.mini-guide {
    background: #ffffff;
    border: 1px solid var(--card-border);
    border-radius: 8px;
    padding: 0.85rem 1rem;
    margin: 0.25rem 0 1rem 0;
}

hr {
    border: none !important;
    border-top: 1px solid var(--card-border) !important;
    margin: 1.4rem 0 !important;
}
</style>
""",
    unsafe_allow_html=True,
)


# ─────────────────────────────────────
# Helpers
# ─────────────────────────────────────
def get_client() -> Anthropic | None:
    api_key = st.secrets.get("ANTHROPIC_API_KEY", os.getenv("ANTHROPIC_API_KEY"))
    return Anthropic(api_key=api_key) if api_key else None


def stream_ai(prompt_text: str, tokens: int = 8000):
    client = get_client()
    if not client:
        yield "❌ ANTHROPIC_API_KEY가 설정되지 않았습니다. Streamlit Secrets 또는 환경변수에 등록해 주세요."
        return

    try:
        with client.messages.stream(
            model=ANTHROPIC_MODEL,
            max_tokens=tokens,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": prompt_text}],
        ) as stream:
            for text in stream.text_stream:
                yield text
    except Exception as exc:
        yield f"\n\n❌ 오류: {exc}"


def generate_to_state(state_key: str, prompt_text: str, label: str, tokens: int = 8000):
    st.markdown(f'<div class="callout"><div class="cl">RUNNING</div>{label}</div>', unsafe_allow_html=True)
    result = st.write_stream(stream_ai(prompt_text, tokens=tokens))
    st.session_state[state_key] = result or ""
    st.rerun()


def chunks() -> list[str]:
    return [
        st.session_state["overview"],
        st.session_state["characters"],
        st.session_state["synopsis"],
        st.session_state["extra_notes"],
    ]


def previous_unit_summary(unit_no: int) -> str:
    if unit_no <= 1:
        return ""
    prev = st.session_state["unit_drafts"].get(unit_no - 1, "")
    if not prev:
        return ""
    return prev[:2500]


def all_outputs_text() -> str:
    parts = []
    if st.session_state["merged_summary"]:
        parts.append("=" * 70 + "\n통합 분석\n" + "=" * 70 + "\n\n" + st.session_state["merged_summary"])
    if st.session_state["gap_report"]:
        parts.append("=" * 70 + "\n부족한 점 진단\n" + "=" * 70 + "\n\n" + st.session_state["gap_report"])
    if st.session_state["reinforced_story"]:
        parts.append("=" * 70 + "\n전체 줄거리 보강\n" + "=" * 70 + "\n\n" + st.session_state["reinforced_story"])
    if st.session_state["unit_plan"]:
        parts.append("=" * 70 + "\n12 Unit 설계\n" + "=" * 70 + "\n\n" + st.session_state["unit_plan"])

    drafts: Dict[int, str] = st.session_state["unit_drafts"]
    for unit_no in sorted(drafts.keys()):
        parts.append(
            "=" * 70
            + f"\nUnit {unit_no:02d} 원고\n"
            + "=" * 70
            + "\n\n"
            + drafts[unit_no]
        )
    return "\n\n\n".join(parts)


def make_docx_bytes(title: str, genre: str, body_text: str) -> bytes:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BLUE JEANS PICTURES")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x20, 0x2A, 0x78)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NOVEL ENGINE")
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x20, 0x2A, 0x78)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"제목: {title or '(제목 미입력)'}  |  장르: {genre}")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x8A, 0x8F, 0xA3)

    doc.add_page_break()

    for line in body_text.split("\n"):
        if line.strip().startswith("="):
            continue
        if line.strip():
            if line.startswith("Unit ") or line in {"통합 분석", "부족한 점 진단", "전체 줄거리 보강", "12 Unit 설계"}:
                h = doc.add_heading(line.strip(), level=1)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x20, 0x2A, 0x78)
            else:
                doc.add_paragraph(line)
        else:
            doc.add_paragraph("")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─────────────────────────────────────
# Header
# ─────────────────────────────────────
st.markdown(
    '<div style="text-align:center;padding:1rem 0 0 0">'
    '<div class="header">B L U E &nbsp; J E A N S &nbsp; P I C T U R E S</div>'
    '<div class="brand-title">NOVEL ENGINE</div>'
    '<div class="sub">C I N E M A T I C &nbsp; · &nbsp; C O M M E R C I A L &nbsp; · &nbsp; S I M P L E &nbsp; · &nbsp; C L I C K</div>'
    '</div>',
    unsafe_allow_html=True,
)

st.markdown(
    '<div class="callout"><div class="cl">HOW TO USE</div>'
    '기획서를 아무도 모를 용어로 자르지 않습니다. 아래 4칸은 의미가 분명합니다. '
    '① 작품 개요 ② 캐릭터 ③ 줄거리/트리트먼트 ④ 추가 메모. '
    '네 칸을 다 채울 필요는 없고, 보통 1~3번만으로도 작동합니다.</div>',
    unsafe_allow_html=True,
)


# ─────────────────────────────────────
# STEP 1
# ─────────────────────────────────────
st.markdown(
    '<div class="section-header">STEP 1 · 기획서 입력 <span class="en">CLEAR INPUT ONLY</span></div>',
    unsafe_allow_html=True,
)

meta_col1, meta_col2 = st.columns([1.4, 1])
with meta_col1:
    st.session_state["title"] = st.text_input(
        "작품 제목",
        value=st.session_state["title"],
        placeholder="예: 머지 앤 어퀴지션",
    )
with meta_col2:
    st.session_state["genre"] = st.selectbox(
        "장르",
        ["스릴러", "드라마", "느와르", "멜로/로맨스", "호러", "액션", "코미디", "SF/판타지"],
        index=["스릴러", "드라마", "느와르", "멜로/로맨스", "호러", "액션", "코미디", "SF/판타지"].index(st.session_state["genre"]),
    )

st.session_state["style_note"] = st.text_input(
    "문체 지향",
    value=st.session_state["style_note"],
    placeholder="예: 시드니 셀던 스타일의 대중소설. 빠르고 선명한 문장. 영화적 장면감.",
)

with st.expander("입력칸 설명 보기", expanded=False):
    st.markdown(
        """
- **작품 개요**: 로그라인, 기획의도, 장르, 톤, 세계관, 차별점
- **캐릭터**: 주인공, 적대자, 조력자, 관계, 욕망, 결핍, 비밀
- **줄거리 / 트리트먼트**: 시작, 중반, 위기, 클라이맥스, 엔딩 방향
- **추가 메모**: 꼭 살릴 장면, 약한 부분, 참고 작품, 조사 메모
        """
    )

col1, col2 = st.columns(2)
with col1:
    st.session_state["overview"] = st.text_area(
        "작품 개요",
        value=st.session_state["overview"],
        height=210,
        placeholder="로그라인 / 기획의도 / 세계관 / 작품의 핵심 질문 / 차별점",
    )
    st.session_state["synopsis"] = st.text_area(
        "줄거리 / 트리트먼트",
        value=st.session_state["synopsis"],
        height=230,
        placeholder="전체 줄거리, 주요 사건, 반전, 엔딩 방향",
    )
with col2:
    st.session_state["characters"] = st.text_area(
        "캐릭터",
        value=st.session_state["characters"],
        height=210,
        placeholder="주인공 / 적대자 / 조력자 / 관계 구조 / 욕망 / 결핍 / 비밀",
    )
    st.session_state["extra_notes"] = st.text_area(
        "추가 메모 (선택)",
        value=st.session_state["extra_notes"],
        height=230,
        placeholder="꼭 살릴 장면, 약한 부분, 참고 문체, 조사 자료, 산업 디테일",
    )


# ─────────────────────────────────────
# STEP 2
# ─────────────────────────────────────
st.markdown(
    '<div class="section-header">STEP 2 · 분석과 설계 <span class="en">ONE CLICK</span></div>',
    unsafe_allow_html=True,
)

st.markdown(
    '<div class="small-meta">프로그램이 자동으로 통합 분석 → 부족한 점 진단 → 전체 줄거리 보강 → 12 Unit 설계를 순서대로 진행합니다.</div>',
    unsafe_allow_html=True,
)

btn_c1, btn_c2, btn_c3, btn_c4 = st.columns(4)

with btn_c1:
    if st.button("기획서 통합 분석", type="primary", use_container_width=True):
        prompt_text = build_intake_merge_prompt(
            title=st.session_state["title"],
            genre=st.session_state["genre"],
            style_note=st.session_state["style_note"],
            chunks=chunks(),
        )
        generate_to_state("merged_summary", prompt_text, "기획서를 하나의 통합 요약으로 분석 중…", tokens=7000)

with btn_c2:
    if st.button("부족한 점 진단", use_container_width=True, disabled=not st.session_state["merged_summary"]):
        prompt_text = build_gap_diagnosis_prompt(
            title=st.session_state["title"],
            genre=st.session_state["genre"],
            merged_summary=st.session_state["merged_summary"],
        )
        generate_to_state("gap_report", prompt_text, "장편화에 부족한 점을 진단 중…", tokens=7000)

with btn_c3:
    if st.button(
        "전체 줄거리 보강",
        use_container_width=True,
        disabled=not (st.session_state["merged_summary"] and st.session_state["gap_report"]),
    ):
        prompt_text = build_story_reinforcement_prompt(
            title=st.session_state["title"],
            genre=st.session_state["genre"],
            merged_summary=st.session_state["merged_summary"],
            gap_report=st.session_state["gap_report"],
        )
        generate_to_state("reinforced_story", prompt_text, "장편소설용 전체 줄거리를 보강 중…", tokens=8000)

with btn_c4:
    if st.button(
        "12 Unit 설계",
        use_container_width=True,
        disabled=not st.session_state["reinforced_story"],
    ):
        prompt_text = build_unit_plan_prompt(
            title=st.session_state["title"],
            genre=st.session_state["genre"],
            reinforced_story=st.session_state["reinforced_story"],
        )
        generate_to_state("unit_plan", prompt_text, "12 Unit 구조를 설계 중…", tokens=9000)

if st.session_state["merged_summary"]:
    st.markdown('<div class="output-card">', unsafe_allow_html=True)
    st.subheader("통합 분석")
    st.markdown(st.session_state["merged_summary"])
    st.markdown('</div>', unsafe_allow_html=True)

if st.session_state["gap_report"]:
    st.markdown('<div class="output-card">', unsafe_allow_html=True)
    st.subheader("부족한 점 진단")
    st.markdown(st.session_state["gap_report"])
    st.markdown('</div>', unsafe_allow_html=True)

if st.session_state["reinforced_story"]:
    st.markdown('<div class="output-card">', unsafe_allow_html=True)
    st.subheader("전체 줄거리 보강")
    st.markdown(st.session_state["reinforced_story"])
    st.markdown('</div>', unsafe_allow_html=True)

if st.session_state["unit_plan"]:
    st.markdown('<div class="output-card">', unsafe_allow_html=True)
    st.subheader("12 Unit 설계")
    st.markdown(st.session_state["unit_plan"])
    st.markdown('</div>', unsafe_allow_html=True)


# ─────────────────────────────────────
# STEP 3
# ─────────────────────────────────────
st.markdown(
    '<div class="section-header">STEP 3 · Unit 원고 작성 <span class="en">WRITE UNIT BY UNIT</span></div>',
    unsafe_allow_html=True,
)

u1, u2, u3 = st.columns([1, 1, 1.2])
with u1:
    st.session_state["selected_unit"] = st.selectbox(
        "작성할 Unit 번호",
        list(range(1, 13)),
        index=int(st.session_state["selected_unit"]) - 1,
    )
with u2:
    st.session_state["rewrite_mode"] = st.selectbox(
        "리라이트 방향",
        [
            "더 상업적으로",
            "더 빠르게",
            "더 감정적으로",
            "더 차갑게",
            "더 스릴러답게",
            "더 여성서사 중심으로",
            "더 영상적으로",
        ],
        index=[
            "더 상업적으로",
            "더 빠르게",
            "더 감정적으로",
            "더 차갑게",
            "더 스릴러답게",
            "더 여성서사 중심으로",
            "더 영상적으로",
        ].index(st.session_state["rewrite_mode"]),
    )
with u3:
    st.markdown(
        '<div class="mini-guide"><strong>Unit 작성 원칙</strong><br>한 번에 장편 전체를 쓰지 않고, Unit 단위로 작성합니다. 먼저 Unit 01부터 시작한 뒤 순서대로 진행하는 방식이 가장 안정적입니다.</div>',
        unsafe_allow_html=True,
    )

unit_btn1, unit_btn2 = st.columns(2)
with unit_btn1:
    if st.button(
        f"Unit {st.session_state['selected_unit']:02d} 원고 생성",
        type="primary",
        use_container_width=True,
        disabled=not st.session_state["unit_plan"],
    ):
        unit_no = int(st.session_state["selected_unit"])
        prompt_text = build_unit_draft_prompt(
            title=st.session_state["title"],
            genre=st.session_state["genre"],
            style_note=st.session_state["style_note"],
            reinforced_story=st.session_state["reinforced_story"],
            unit_plan=st.session_state["unit_plan"],
            unit_number=unit_no,
            previous_unit_summary=previous_unit_summary(unit_no),
        )
        st.markdown(
            f'<div class="callout"><div class="cl">WRITING</div>Unit {unit_no:02d} 원고를 작성 중…</div>',
            unsafe_allow_html=True,
        )
        result = st.write_stream(stream_ai(prompt_text, tokens=9000))
        drafts = dict(st.session_state["unit_drafts"])
        drafts[unit_no] = result or ""
        st.session_state["unit_drafts"] = drafts
        st.rerun()

with unit_btn2:
    selected_text = st.session_state["unit_drafts"].get(int(st.session_state["selected_unit"]), "")
    if st.button(
        f"Unit {st.session_state['selected_unit']:02d} 다시 쓰기",
        use_container_width=True,
        disabled=not bool(selected_text),
    ):
        unit_no = int(st.session_state["selected_unit"])
        prompt_text = build_rewrite_prompt(
            mode=st.session_state["rewrite_mode"],
            text=st.session_state["unit_drafts"].get(unit_no, ""),
        )
        st.markdown(
            f'<div class="callout"><div class="cl">REWRITE</div>Unit {unit_no:02d} 원고를 다시 쓰는 중…</div>',
            unsafe_allow_html=True,
        )
        result = st.write_stream(stream_ai(prompt_text, tokens=9000))
        drafts = dict(st.session_state["unit_drafts"])
        drafts[unit_no] = result or ""
        st.session_state["unit_drafts"] = drafts
        st.rerun()

if st.session_state["unit_drafts"]:
    for unit_no in sorted(st.session_state["unit_drafts"].keys()):
        with st.expander(f"Unit {unit_no:02d} 원고 보기", expanded=(unit_no == int(st.session_state["selected_unit"]))):
            st.markdown(st.session_state["unit_drafts"][unit_no])


# ─────────────────────────────────────
# STEP 4
# ─────────────────────────────────────
st.markdown(
    '<div class="section-header">STEP 4 · 저장 <span class="en">DOWNLOAD</span></div>',
    unsafe_allow_html=True,
)

compiled_text = all_outputs_text()
if compiled_text.strip():
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    filename_base = (st.session_state["title"] or "novel_project").replace(" ", "_")

    dl1, dl2 = st.columns(2)
    with dl1:
        st.download_button(
            label="TXT 저장",
            data=compiled_text,
            file_name=f"{filename_base}_{ts}.txt",
            mime="text/plain",
            use_container_width=True,
        )
    with dl2:
        try:
            docx_bytes = make_docx_bytes(st.session_state["title"], st.session_state["genre"], compiled_text)
            st.download_button(
                label="DOCX 저장",
                data=docx_bytes,
                file_name=f"{filename_base}_{ts}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except ImportError:
            st.caption("DOCX 저장을 위해 python-docx 설치가 필요합니다.")
else:
    st.caption("아직 저장할 결과가 없습니다. 먼저 분석 또는 Unit 원고를 생성해 주세요.")


# ─────────────────────────────────────
# Footer
# ─────────────────────────────────────
st.markdown("---")
fc1, fc2 = st.columns([3, 1])
with fc2:
    if st.button("전체 초기화", use_container_width=True):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

st.caption("© 2026 BLUE JEANS PICTURES · Novel Engine v1.1")
